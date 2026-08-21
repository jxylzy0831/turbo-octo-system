from pathlib import Path
import re
import argparse
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
COVER = ROOT / "assets" / "阿奇封面-v1.png"
OUT = ROOT / "output"
ORIGINALS = ROOT / "source" / "originals"

def set_font(run, name="宋体", size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_page_number(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r._r.addnext(fld)
    set_font(r, "Times New Roman", 9)

def new_decimal_numbering(doc):
    numbering = doc.part.numbering_part.element
    abs_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abs_id = max(abs_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abs_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "decimal")
    lvl.append(fmt)
    text_el = OxmlElement("w:lvlText")
    text_el.set(qn("w:val"), "%1.")
    lvl.append(text_el)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "space")
    lvl.append(suff)
    ppr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    ppr.append(ind)
    lvl.append(ppr)
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abs_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id

def apply_decimal_numbering(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    numpr.append(ilvl)
    numpr.append(numid)
    ppr.append(numpr)

def setup(doc):
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(14.8), Cm(21)
    sec.top_margin, sec.bottom_margin = Cm(1.8), Cm(1.7)
    sec.left_margin, sec.right_margin = Cm(1.75), Cm(1.75)
    text_direction = sec._sectPr.find(qn("w:textDirection"))
    if text_direction is None:
        text_direction = OxmlElement("w:textDirection")
        sec._sectPr.append(text_direction)
    text_direction.set(qn("w:val"), "lrTb")
    columns = sec._sectPr.find(qn("w:cols"))
    if columns is None:
        columns = OxmlElement("w:cols")
        sec._sectPr.append(columns)
    columns.set(qn("w:num"), "1")
    add_page_number(sec)
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.45
    normal.paragraph_format.space_after = Pt(5.5)
    normal.paragraph_format.first_line_indent = Cm(0.74)
    for name, size in [("Title", 26), ("Heading 1", 20), ("Heading 2", 15), ("Heading 3", 12)]:
        try:
            s = doc.styles[name]
        except KeyError:
            s = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            s.base_style = normal
        s.font.name = "宋体"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        s.font.size = Pt(size)
        s.font.bold = True
    try:
        doc.styles["List Bullet"]
    except KeyError:
        s = doc.styles.add_style("List Bullet", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = normal
        s.paragraph_format.left_indent = Cm(0.74)
        s.paragraph_format.first_line_indent = Cm(-0.37)

def clear_document_body(doc):
    """Clear the copied rolebook body while retaining its package, theme and styles."""
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)

def cover(doc, volume):
    sec = doc.sections[0]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(COVER), width=Cm(9.25))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("流氓叙事")
    set_font(r, "宋体", 24, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"阿奇 · {volume}本")
    set_font(r, "宋体", 16, True, (132, 38, 46))
    doc.add_page_break()

def add_markdown(doc, text):
    active_num_id = None
    for raw in text.splitlines():
        line = raw.strip()
        if not line or line == "---":
            if line == "---":
                p = doc.add_paragraph("◆")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(12)
            if line == "---":
                active_num_id = None
            continue
        if line.startswith("# "):
            continue
        if line.startswith("## "):
            title = line[3:]
            if title.startswith("阿奇 ·"):
                continue
            if title.startswith("演员须知") or title.startswith("最终演绎") or title.startswith("演员伴手礼") or title.startswith("阿奇的绝笔") or title.startswith("A本收束"):
                doc.add_page_break()
            p = doc.add_paragraph(style="Heading 1")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(16)
            p.add_run(title)
            continue
        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(10)
            p.add_run(line[4:])
            continue
        if line.startswith("> "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            r = p.add_run(line[2:])
            set_font(r, "楷体", 12, False, (120, 35, 42))
            continue
        m = re.match(r"^(\d+)\.\s+(.*)", line)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0.74)
            p.paragraph_format.hanging_indent = Cm(0.37)
            p.add_run(f"{m.group(1)}. {m.group(2)}")
            continue
        active_num_id = None
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.first_line_indent = Cm(0)
            p.add_run(line[2:])
            continue
        p = doc.add_paragraph()
        if line.startswith("你需要做到") or line.startswith("你可以说") or line.startswith("你不可以说") or line.startswith("从此处开始") or line.startswith("推荐末句") or line.startswith("若现场"):
            p.paragraph_format.first_line_indent = Cm(0)
            r = p.add_run(line)
            set_font(r, "宋体", 10.5, True)
        else:
            p.add_run(line)

def build(src, target=None):
    volume = re.search(r"阿奇([ABC])本", src.stem).group(1)
    # AGENTS.md requires the final file to derive from an original rolebook copy.
    # Loading the matching 黛利拉 volume preserves the original package/theme/style
    # lineage; only the copied document body is cleared. The source file is never saved.
    template = ORIGINALS / f"黛利拉{volume}本.docx"
    doc = Document(template)
    clear_document_body(doc)
    setup(doc)
    cover(doc, volume)
    text = src.read_text(encoding="utf-8")
    add_markdown(doc, text)
    props = doc.core_properties
    props.title = f"《流氓叙事》阿奇{volume}本"
    props.subject = "NPC演员角色本与演后收藏本"
    props.author = "《流氓叙事》阿奇性转改写工程"
    target = Path(target) if target else OUT / f"阿奇{volume}本.docx"
    doc.save(target)
    print(target)

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--source")
    parser.add_argument("--output")
    args = parser.parse_args()
    OUT.mkdir(exist_ok=True)
    if args.source:
        build(Path(args.source), args.output)
    else:
        versions = {
            "A": ROOT / "drafts" / "阿奇A本-v3.md",
            "B": ROOT / "drafts" / "阿奇B本-v2.md",
            "C": ROOT / "drafts" / "阿奇C本-v2.md",
        }
        for volume, src in versions.items():
            build(src, OUT / f"阿奇{volume}本-定稿.docx")
