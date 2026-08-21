from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

import build_archie_docx as base

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output"
ORIGINALS = ROOT / "source" / "originals"

VOLUMES = [
    ("A", ROOT / "drafts" / "阿奇A本-v3.md"),
    ("B", ROOT / "drafts" / "阿奇B本-v2.md"),
    ("C", ROOT / "drafts" / "阿奇C本-v2.md"),
]


def combined_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(base.COVER), width=Cm(9.25))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("流氓叙事")
    base.set_font(r, "宋体", 24, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("阿奇 · 合订本")
    base.set_font(r, "宋体", 16, True, (132, 38, 46))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("A本 · B本 · C本")
    base.set_font(r, "宋体", 11, False, (120, 35, 42))
    doc.add_page_break()


def volume_marker(doc, volume):
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(16)
    p.add_run(f"阿奇 · {volume}本")


def build_combined(target=None):
    template = ORIGINALS / "黛利拉A本.docx"
    doc = Document(template)
    base.clear_document_body(doc)
    base.setup(doc)
    combined_cover(doc)
    for index, (volume, src) in enumerate(VOLUMES):
        if index:
            doc.add_page_break()
        volume_marker(doc, volume)
        base.add_markdown(doc, src.read_text(encoding="utf-8"))
    props = doc.core_properties
    props.title = "《流氓叙事》阿奇 A/B/C 合订本"
    props.subject = "NPC演员角色本与演后收藏本"
    props.author = "《流氓叙事》阿奇性转改写工程"
    target = Path(target) if target else OUT / "阿奇ABC本-合订本.docx"
    doc.save(target)
    print(target)


if __name__ == "__main__":
    OUT.mkdir(exist_ok=True)
    build_combined()
